import csv
import json
import re
import unicodedata
from pathlib import Path

# File dữ liệu mặc định: <project>/data/accounts.json
DEFAULT_DATA_FILE = Path(__file__).resolve().parent.parent / "data" / "accounts.json"

REQUIRED_FIELDS = {"username", "email", "password"}

# Tên cột/tiêu đề có thể gặp -> trường chuẩn. So khớp sau khi bỏ dấu, viết
# thường. Nhờ vậy "Mật khẩu", "MAT KHAU", "password" đều nhận ra như nhau.
FIELD_ALIASES = {
    "username": [
        "username", "user", "user name", "login", "account",
        "tai khoan", "ten tai khoan", "ten dang nhap", "tendangnhap",
    ],
    "email": [
        "email", "e mail", "mail", "thu dien tu", "dia chi email",
    ],
    "password": [
        "password", "pass", "pwd", "mat khau", "matkhau",
    ],
    "full_name": [
        "full name", "fullname", "ho ten", "ho va ten", "ho va ten day du",
        "ten day du", "ten", "name",
    ],
}


def _normalize(text: str) -> str:
    """Chuẩn hóa tiêu đề: bỏ dấu, đ->d, viết thường, gộp khoảng trắng."""
    if text is None:
        return ""
    s = str(text).strip().lower().replace("đ", "d")
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")  # bỏ dấu
    s = re.sub(r"[\s_\-]+", " ", s).strip()
    return s


# Bảng tra nhanh: tiêu đề chuẩn hóa -> trường chuẩn
_ALIAS_LOOKUP = {}
for _field, _names in FIELD_ALIASES.items():
    for _n in _names:
        _ALIAS_LOOKUP[_normalize(_n)] = _field


def _map_header(header: str):
    """Trả về tên trường chuẩn cho một tiêu đề cột, hoặc None nếu không khớp."""
    return _ALIAS_LOOKUP.get(_normalize(header))


def _build_account(record: dict) -> dict:
    """Chuyển một bản ghi thô (tiêu đề gốc -> giá trị) thành tài khoản chuẩn."""
    acc = {}
    for header, value in record.items():
        field = _map_header(header)
        if field and value not in (None, ""):
            acc[field] = str(value).strip()
    return acc


# ----------------------- Các bộ đọc theo định dạng -----------------------
def _read_json(path: Path) -> list[dict]:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, list):
        raise ValueError("File JSON phải là một mảng các tài khoản.")
    return [_build_account(row) for row in data]


def _read_csv(path: Path) -> list[dict]:
    with open(path, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        return [_build_account(row) for row in reader]


def _read_xlsx(path: Path) -> list[dict]:
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    rows = ws.iter_rows(values_only=True)
    try:
        headers = list(next(rows))
    except StopIteration:
        return []
    accounts = []
    for row in rows:
        if row is None or all(c is None for c in row):
            continue
        record = {headers[i]: row[i] for i in range(len(headers)) if i < len(row)}
        accounts.append(_build_account(record))
    return accounts


def _read_docx(path: Path) -> list[dict]:
    from docx import Document
    doc = Document(path)
    accounts = []

    # 1) Ưu tiên đọc từ BẢNG (hàng đầu = tiêu đề)
    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue
        headers = [c.text for c in rows[0].cells]
        for row in rows[1:]:
            cells = [c.text for c in row.cells]
            if not any(cell.strip() for cell in cells):
                continue
            record = {headers[i]: cells[i] for i in range(min(len(headers), len(cells)))}
            accounts.append(_build_account(record))

    # 2) Nếu không có bảng, đọc theo dòng "Nhãn: giá trị", mỗi tài khoản cách
    #    nhau bằng dòng trống.
    if not accounts:
        current = {}
        for para in doc.paragraphs:
            line = para.text.strip()
            if not line:
                if current:
                    accounts.append(_build_account(current))
                    current = {}
                continue
            if ":" in line:
                key, val = line.split(":", 1)
                current[key.strip()] = val.strip()
        if current:
            accounts.append(_build_account(current))

    return accounts


# Bảng phân loại theo đuôi file
_READERS = {
    ".json": _read_json,
    ".csv": _read_csv,
    ".xlsx": _read_xlsx,
    ".xlsm": _read_xlsx,
    ".docx": _read_docx,
}


def load_accounts(file_path: str = None) -> list[dict]:
    """Đọc danh sách tài khoản từ file và trả về list các dict chuẩn.

    Hỗ trợ: .json, .csv, .xlsx, .docx
    Tự nhận diện loại file theo đuôi và tự map cột (kể cả tiêu đề tiếng Việt
    có dấu) về các trường: username, email, password, full_name.
    """
    path = Path(file_path) if file_path else DEFAULT_DATA_FILE

    if not path.exists():
        raise FileNotFoundError(f"Không tìm thấy file dữ liệu: {path}")

    reader = _READERS.get(path.suffix.lower())
    if reader is None:
        supported = ", ".join(sorted(_READERS))
        raise ValueError(
            f"Không hỗ trợ định dạng '{path.suffix}'. Định dạng hỗ trợ: {supported}"
        )

    accounts = reader(path)

    # Kiểm tra các trường bắt buộc
    valid = []
    for i, acc in enumerate(accounts, start=1):
        missing = REQUIRED_FIELDS - acc.keys()
        if missing:
            raise ValueError(
                f"Tài khoản #{i} thiếu trường bắt buộc: {sorted(missing)} "
                f"(đọc được: {acc})"
            )
        valid.append(acc)

    if not valid:
        raise ValueError(f"Không đọc được tài khoản nào từ file: {path}")

    return valid
